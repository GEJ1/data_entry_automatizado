"""
Render a DOCX del MISMO lote falso que el PDF. Los datos viven en `datos_fake.py`.

Existe para demostrar el punto central del proyecto: cambiar el formato de
entrada NO deberia obligar a tocar el resto del pipeline. Este DOCX tiene el
mismo contenido que el PDF, asi que comparte el mismo ground truth: si el
extractor de PDF y el de DOCX cumplen el contrato, los dos tienen que producir
exactamente las mismas FichaCruda.

Uso:
    python demo/generar_docx_fake.py
    python demo/generar_docx_fake.py --clientes 300 --salida data/entrada/lote.docx

Requiere: python-docx. Los datos son 100% inventados.
"""
import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from datos_fake import COLS_ALE, COLS_REF, armar_lote, manifiesto, verdad


def _tabla(doc, columnas, filas_por_grupo):
    """
    Tabla con encabezado + (opcional) filas de subgrupo combinadas.

    `filas_por_grupo` es una lista de (titulo | None, filas).
    """
    tabla = doc.add_table(rows=1, cols=len(columnas))
    tabla.style = "Table Grid"
    for celda, texto in zip(tabla.rows[0].cells, columnas):
        celda.text = texto
        for p in celda.paragraphs:
            for r in p.runs:
                r.bold = True

    for titulo, filas in filas_por_grupo:
        if titulo is not None:
            fila = tabla.add_row()
            # Celda COMBINADA, igual que el reporte real. Ojo con esto al
            # extraer: python-docx devuelve el mismo texto repetido en las 12
            # columnas (el PDF, en cambio, lo devuelve truncado). Misma
            # estructura, distintas mañas segun el formato.
            combinada = fila.cells[0].merge(fila.cells[-1])
            combinada.text = titulo
            for p in combinada.paragraphs:
                for r in p.runs:
                    r.bold = True
        for f in filas:
            celdas = tabla.add_row().cells
            for celda, col in zip(celdas, columnas):
                celda.text = str(f[col])
    return tabla


def generar(salida, n_clientes):
    clientes = armar_lote(n_clientes)
    doc = Document()
    doc.styles["Normal"].font.size = Pt(8)

    for i, (_, cli) in enumerate(clientes):
        p = doc.add_paragraph()
        r = p.add_run(f"Cliente: {cli['nombre']} (CUIT: {cli['cuit']})")
        r.bold = True
        r.font.size = Pt(12)
        doc.add_paragraph(
            f"Fecha de Emisión: {cli['emision'].strftime('%d/%m/%Y')} 19:00")

        doc.add_paragraph("1. Solicitudes de Referencia")
        if cli["referencias"]:
            _tabla(doc, COLS_REF,
                   [(sg["titulo"], sg["filas"]) for sg in cli["referencias"]])
        else:
            doc.add_paragraph("No hay solicitudes de referencia para este cliente.")

        doc.add_paragraph("2. Alertas / Denuncias")
        if cli["alertas"]:
            _tabla(doc, COLS_ALE, [(None, cli["alertas"])])
        else:
            doc.add_paragraph("No hay alertas ni denuncias para este cliente.")

        if i < len(clientes) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    Path(salida).parent.mkdir(parents=True, exist_ok=True)
    doc.save(salida)

    ruta_verdad = Path(salida).with_suffix(".verdad.json")
    ruta_verdad.write_text(
        json.dumps(verdad(clientes), ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"\nDOCX generado:  {salida}  ({len(clientes)} clientes)")
    print(f"Ground truth:   {ruta_verdad}  (el MISMO contenido que el PDF)\n")
    print(manifiesto(clientes))


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--clientes", type=int, default=15)
    ap.add_argument("--salida", default="data/entrada/lote17.docx")
    args = ap.parse_args()
    generar(args.salida, args.clientes)
