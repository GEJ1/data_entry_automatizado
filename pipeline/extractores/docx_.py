"""
Extractor de DOCX con python-docx.

Se llama `docx_.py` con guion bajo al final para no taparle el nombre a la
libreria `docx` cuando se importa desde adentro del paquete.

Comparalo con `pdf_plumber.py`: por dentro no se parecen en nada (uno mide
posiciones en la hoja, este recorre el XML del documento), y por fuera son
intercambiables. Eso es todo lo que el contrato tiene que garantizar.

Las mañas propias de este formato:

1. NO HAY PAGINAS. Un DOCX es un flujo continuo, asi que el problema del
   "cliente que se derrama" simplemente no existe. El corte por encabezado
   `Cliente: ... (CUIT: ...)` sigue siendo el mismo, igual que en PDF.

2. LAS CELDAS COMBINADAS SE REPITEN. Donde el PDF devolvia el titulo del
   subgrupo TRUNCADO, python-docx lo devuelve REPETIDO en las 12 columnas.
   Distinta maña, misma solucion de fondo: detectar que la fila tiene un solo
   valor distinto y tratarla como titulo, no como dato.

3. HAY QUE RECORRER EN ORDEN. `doc.paragraphs` y `doc.tables` son dos listas
   separadas y se pierde el orden entre ellas; para saber que tabla va con que
   cliente hay que caminar el cuerpo del documento tal como esta escrito.
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterator

import sys

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from pipeline.contratos import FichaCruda, Fila
from pipeline.extractores.comun import (
    PREFIJO_SUBGRUPO, RE_CLIENTE, RE_EMISION, limpiar, seccion_de,
)


def _bloques(doc):
    """
    Parrafos y tablas EN EL ORDEN en que aparecen en el documento.

    python-docx expone doc.paragraphs y doc.tables por separado, y asi se pierde
    la relacion entre un encabezado de cliente y las tablas que lo siguen. La
    unica forma de conservarla es recorrer el cuerpo del XML a mano.
    """
    for hijo in doc.element.body.iterchildren():
        if hijo.tag == qn("w:p"):
            yield Paragraph(hijo, doc)
        elif hijo.tag == qn("w:tbl"):
            yield Table(hijo, doc)


def _filas(tabla: Table) -> list[list[str]]:
    return [[limpiar(c.text) for c in fila.cells] for fila in tabla.rows]


class ExtractorDOCX:
    """Extractor para los reportes en DOCX. Ver contratos.Extractor."""

    formatos = (".docx",)
    nombre = "docx"

    def extraer(self, ruta: Path) -> Iterator[FichaCruda]:
        ruta = Path(ruta)
        doc = Document(str(ruta))

        ficha: FichaCruda | None = None
        solicitud = ""
        n_ficha = 0

        for bloque in _bloques(doc):
            if isinstance(bloque, Paragraph):
                texto = limpiar(bloque.text)
                if not texto:
                    continue

                m = RE_CLIENTE.match(texto)
                if m:
                    if ficha is not None:
                        yield ficha
                    n_ficha += 1
                    ficha = FichaCruda(
                        cabecera={"Cliente": m.group("nombre"),
                                  "CUIT": m.group("cuit")},
                        # Ambas secciones siempre presentes: "no hay tabla" y
                        # "tabla vacia" son lo mismo para quien consume esto.
                        tablas={"referencias": [], "alertas": []},
                        origen=f"{ruta.name} ficha {n_ficha}")
                    solicitud = ""
                    continue

                m = RE_EMISION.match(texto)
                if m and ficha is not None:
                    ficha.cabecera["Fecha de Emisión"] = m.group("valor")

            elif isinstance(bloque, Table) and ficha is not None:
                solicitud = self._volcar_tabla(bloque, ficha, solicitud, ruta)

        if ficha is not None:
            yield ficha

    # -- internos -----------------------------------------------------------

    @staticmethod
    def _volcar_tabla(tabla: Table, ficha: FichaCruda, solicitud: str,
                      ruta: Path) -> str:
        """Vuelca una tabla en la ficha. Devuelve el subgrupo vigente al terminar."""
        filas = _filas(tabla)
        if not filas:
            return solicitud

        encabezados = filas[0]
        seccion = seccion_de(encabezados)
        if seccion is None:
            # Tabla que no reconocemos: se avisa fuerte en vez de tragarsela.
            print(f"AVISO: tabla desconocida en {ruta.name}, "
                  f"encabezados={encabezados}", file=sys.stderr)
            return solicitud

        for celdas in filas[1:]:
            distintos = {c for c in celdas if c}
            if not distintos:
                continue  # fila totalmente vacia

            # Fila de subgrupo: la celda combinada devuelve el MISMO texto en
            # todas las columnas, asi que hay un solo valor distinto.
            if (seccion == "referencias" and len(distintos) == 1
                    and next(iter(distintos)).startswith(PREFIJO_SUBGRUPO)):
                solicitud = next(iter(distintos))
                continue

            fila: Fila = dict(zip(encabezados, celdas))
            if seccion == "referencias":
                # La solicitud es un titulo que abarca varias filas en el
                # documento; aca se baja a cada fila porque despues forma
                # parte de la clave que identifica el registro.
                fila = {"Solicitud": solicitud, **fila}
            ficha.tablas[seccion].append(fila)

        return solicitud
